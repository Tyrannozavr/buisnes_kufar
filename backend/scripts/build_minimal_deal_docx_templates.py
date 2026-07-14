"""
Пересобирает bill*.docx и supply_contract.docx как минимальные шаблоны docxtpl.

Старые файлы из Word часто ломают Jinja (плейсхолдеры разбиты на несколько w:r).
Запуск из корня backend: python scripts/build_minimal_deal_docx_templates.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

CONTRACT_FONT_NAME = "Times New Roman"
CONTRACT_FONT_SIZE = Pt(12)
CONTRACT_RIGHT_COLUMN_TAB = Cm(9)


def _set_run_font(run, *, bold: bool = False) -> None:
	run.font.name = CONTRACT_FONT_NAME
	run.font.size = CONTRACT_FONT_SIZE
	run.bold = bold
	r_pr = run._element.get_or_add_rPr()
	r_fonts = OxmlElement("w:rFonts")
	r_fonts.set(qn("w:ascii"), CONTRACT_FONT_NAME)
	r_fonts.set(qn("w:hAnsi"), CONTRACT_FONT_NAME)
	r_fonts.set(qn("w:cs"), CONTRACT_FONT_NAME)
	r_fonts.set(qn("w:eastAsia"), CONTRACT_FONT_NAME)
	r_pr.insert(0, r_fonts)


def _apply_contract_document_style(doc: Document) -> None:
	normal = doc.styles["Normal"]
	normal.font.name = CONTRACT_FONT_NAME
	normal.font.size = CONTRACT_FONT_SIZE
	normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	normal.paragraph_format.first_line_indent = Cm(1.25)

	for section in doc.sections:
		section.top_margin = Cm(2)
		section.bottom_margin = Cm(2)
		section.left_margin = Cm(2)
		section.right_margin = Cm(1)


def _para(doc: Document, text: str) -> None:
	p = doc.add_paragraph()
	run = p.add_run(text)
	_set_run_font(run)


def _contract_title(doc: Document, text: str) -> None:
	p = doc.add_paragraph()
	p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	p.paragraph_format.first_line_indent = Cm(0)
	run = p.add_run(text)
	_set_run_font(run, bold=True)


def _contract_centered(doc: Document, text: str) -> None:
	p = doc.add_paragraph()
	p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	p.paragraph_format.first_line_indent = Cm(0)
	run = p.add_run(text)
	_set_run_font(run)


def _contract_body_para(doc: Document, text: str) -> None:
	p = doc.add_paragraph()
	p.paragraph_format.first_line_indent = Cm(1.25)
	p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	run = p.add_run(text)
	_set_run_font(run)


def _contract_page_break(doc: Document) -> None:
	p = doc.add_paragraph()
	p.paragraph_format.first_line_indent = Cm(0)
	run = p.add_run()
	run.add_break(WD_BREAK.PAGE)


def _set_paragraph_keep(paragraph, *, together: bool = True, with_next: bool = False) -> None:
	paragraph.paragraph_format.keep_together = together
	paragraph.paragraph_format.keep_with_next = with_next


def _add_two_column_lines(
	doc: Document,
	rows: list[tuple[str, str]],
	*,
	bold_header: bool = False,
	page_break_before: bool = False,
	keep_rows_together: bool = False,
) -> None:
	"""Две колонки без таблицы — табуляция, как в превью (без видимой сетки Word)."""
	if page_break_before:
		_contract_page_break(doc)

	last_row_idx = len(rows) - 1
	for row_idx, (left, right) in enumerate(rows):
		p = doc.add_paragraph()
		p.paragraph_format.first_line_indent = Cm(0)
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
		p.paragraph_format.tab_stops.add_tab_stop(CONTRACT_RIGHT_COLUMN_TAB, WD_TAB_ALIGNMENT.LEFT)
		keep_with_next = keep_rows_together and row_idx < last_row_idx
		_set_paragraph_keep(p, together=True, with_next=keep_with_next)

		is_header = bold_header and row_idx == 0
		left_run = p.add_run(left)
		_set_run_font(left_run, bold=is_header)
		p.add_run("\t")
		right_run = p.add_run(right)
		_set_run_font(right_run, bold=is_header)

	doc.add_paragraph()


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
	cell.text = ""
	p = cell.paragraphs[0]
	p.paragraph_format.first_line_indent = Cm(0)
	run = p.add_run(text)
	_set_run_font(run, bold=bold)


def _add_bank_header_table(doc: Document) -> None:
	"""Шапка с реквизитами банка/получателя — как на бланке счёта."""
	table = doc.add_table(rows=4, cols=4)
	table.style = "Table Grid"
	_set_cell_text(table.rows[0].cells[0], "{{ seller_company.bank_name or '' }}")
	table.rows[0].cells[0].merge(table.rows[0].cells[1])
	_set_cell_text(table.rows[0].cells[2], "БИК", bold=True)
	_set_cell_text(table.rows[0].cells[3], "{{ seller_company.bic or '' }}")

	_set_cell_text(table.rows[1].cells[0], "Банк получателя")
	table.rows[1].cells[0].merge(table.rows[1].cells[1])
	_set_cell_text(table.rows[1].cells[2], "Сч. №", bold=True)
	_set_cell_text(table.rows[1].cells[3], "{{ seller_company.correspondent_bank_account or '' }}")

	_set_cell_text(table.rows[2].cells[0], "ИНН {{ seller_company.inn or '' }}")
	_set_cell_text(table.rows[2].cells[1], "КПП {{ seller_company.kpp or '' }}")
	_set_cell_text(table.rows[2].cells[2], "Сч. №", bold=True)
	_set_cell_text(table.rows[2].cells[3], "{{ seller_company.account_number or '' }}")

	_set_cell_text(
		table.rows[3].cells[0],
		"{{ seller_company.company_type or '' }} {{ seller_company.company_name or '' }}\nПолучатель",
	)
	table.rows[3].cells[0].merge(table.rows[3].cells[3])
	doc.add_paragraph()


def _add_items_table(doc: Document) -> None:
	"""
	Таблица позиций с циклом docxtpl {%tr for %}.

	for/endfor обязаны быть в ОТДЕЛЬНЫХ строках таблицы: в docxtpl 0.18 regex
	жадный и при обоих тегах в одной <w:tr> оставляет только {% endfor %}.
	"""
	table = doc.add_table(rows=4, cols=7)
	table.style = "Table Grid"
	headers = ("№", "Название продукта", "Артикул", "Кол-во", "Ед. изм.", "Цена", "Сумма")
	for i, header in enumerate(headers):
		_set_cell_text(table.rows[0].cells[i], header, bold=True)

	# Строка только с открытием цикла (целиком заменяется на {% for ... %})
	_set_cell_text(table.rows[1].cells[0], "{%tr for item in items %}")

	row = table.rows[2].cells
	_set_cell_text(row[0], "{{ loop.index }}")
	_set_cell_text(row[1], "{{ item.product_name }}")
	_set_cell_text(row[2], "{{ item.product_article }}")
	_set_cell_text(row[3], "{{ item.quantity }}")
	_set_cell_text(row[4], "{{ item.unit_of_measurement }}")
	_set_cell_text(row[5], "{{ item.price }}")
	_set_cell_text(row[6], "{{ item.amount }}")

	# Строка только с закрытием цикла
	_set_cell_text(table.rows[3].cells[0], "{%tr endfor %}")
	doc.add_paragraph()


def _add_bill_common_body(doc: Document, *, title_jinja: str) -> None:
	_add_bank_header_table(doc)
	_contract_title(doc, title_jinja)
	_para(doc, "Поставщик (исполнитель): {{ seller_party_line }}")
	_para(doc, "Покупатель (заказчик): {{ buyer_party_line }}")
	_para(doc, "{% if bill.reason %}Основание: {{ bill.reason }}{% endif %}")
	_add_items_table(doc)
	_para(doc, "Итого: {{ total_amount_excl_vat }}")
	_para(
		doc,
		"В том числе НДС: {% if bill.show_vat_row %}{{ bill.vat_amount_display }}{% endif %}",
	)
	_para(doc, "Всего к оплате: {{ total_amount }}")
	_para(doc, "Всего наименований: {{ items_count }}, на сумму: {{ total_amount }} руб.")
	_para(doc, "{{ total_word }}")


def _add_officials_signature(doc: Document) -> None:
	# for/endfor в разных абзацах — иначе {%p ...%} в одном <w:p> сжирает endfor
	_para(doc, "{%p for o in bill.officials %}")
	_para(doc, "{{ o.position or '' }} _______________ /{{ o.full_name or '' }}/")
	_para(doc, "{%p endfor %}")
	_para(doc, "_______________ /(должность, подпись, ФИО)/")

def write_bill_payment(path: Path) -> None:
	doc = Document()
	_add_bill_common_body(
		doc,
		title_jinja="Счет на оплату № {{ bill_number }} от {{ bill_date_fmt }} г.",
	)
	_para(doc, "{% if bill.payment_validity_text %}{{ bill.payment_validity_text }}{% endif %}")
	_para(doc, "{% if bill.additional_info %}{{ bill.additional_info }}{% endif %}")
	_add_officials_signature(doc)
	path.parent.mkdir(parents=True, exist_ok=True)
	doc.save(str(path))
	print("Wrote", path)


def write_bill_contract(path: Path) -> None:
	doc = Document()
	_add_bill_common_body(
		doc,
		title_jinja="Счет-договор № {{ bill_number }} от {{ bill_date_fmt }} г.",
	)
	_para(doc, "{% if bill.contract_terms_text %}{{ bill.contract_terms_text }}{% endif %}")
	_add_two_column_lines(
		doc,
		[
			(
				"{% if show_supplier_details %}ПОСТАВЩИК:{% endif %}",
				"{% if show_buyer_details %}ПОКУПАТЕЛЬ:{% endif %}",
			),
			(
				"{% if show_supplier_details %}{{ seller_company.company_name or '' }}{% endif %}",
				"{% if show_buyer_details %}{{ buyer_company.company_name or '' }}{% endif %}",
			),
			(
				"{% if show_supplier_details %}{{ seller_company.index or '' }} {{ seller_company.legal_address or '' }}{% endif %}",
				"{% if show_buyer_details %}{{ buyer_company.index or '' }} {{ buyer_company.legal_address or '' }}{% endif %}",
			),
			(
				"{% if show_supplier_details %}ИНН: {{ seller_company.inn or '' }} КПП: {{ seller_company.kpp or '' }}{% endif %}",
				"{% if show_buyer_details %}ИНН: {{ buyer_company.inn or '' }} КПП: {{ buyer_company.kpp or '' }}{% endif %}",
			),
			(
				"{% if show_supplier_details %}Рас/счет №: {{ seller_company.account_number or '' }}{% endif %}",
				"{% if show_buyer_details %}Рас/счет №: {{ buyer_company.account_number or '' }}{% endif %}",
			),
			(
				"{% if show_supplier_details %}Корр/счет: {{ seller_company.correspondent_bank_account or '' }}{% endif %}",
				"{% if show_buyer_details %}Корр/счет: {{ buyer_company.correspondent_bank_account or '' }}{% endif %}",
			),
			(
				"{% if show_supplier_details %}Банк: {{ seller_company.bank_name or '' }}{% endif %}",
				"{% if show_buyer_details %}Банк: {{ buyer_company.bank_name or '' }}{% endif %}",
			),
			(
				"{% if show_supplier_details %}БИК: {{ seller_company.bic or '' }}{% endif %}",
				"{% if show_buyer_details %}БИК: {{ buyer_company.bic or '' }}{% endif %}",
			),
		],
		bold_header=True,
		keep_rows_together=True,
	)
	_add_officials_signature(doc)
	path.parent.mkdir(parents=True, exist_ok=True)
	doc.save(str(path))
	print("Wrote", path)


def write_bill_offer(path: Path) -> None:
	doc = Document()
	_add_bill_common_body(
		doc,
		title_jinja="Счет-оферта № {{ bill_number }} от {{ bill_date_fmt }} г.",
	)
	_para(
		doc,
		"{% if bill.additional_info_offer %}{{ bill.additional_info_offer }}{% endif %}",
	)
	_para(doc, "{% if bill.contract_terms_text %}Условия счета-оферты:{% endif %}")
	_para(doc, "{% if bill.contract_terms_text %}{{ bill.contract_terms_text }}{% endif %}")
	_add_officials_signature(doc)
	path.parent.mkdir(parents=True, exist_ok=True)
	doc.save(str(path))
	print("Wrote", path)


def write_bill_variant(path: Path, title: str) -> None:
	"""Обратная совместимость вызова из старых скриптов."""
	if "оферт" in title.lower() or "offer" in path.name:
		write_bill_offer(path)
	else:
		write_bill_contract(path)


def write_supply_contract(path: Path) -> None:
	doc = Document()
	_apply_contract_document_style(doc)
	_contract_title(doc, "ДОГОВОР ПОСТАВКИ № {{ supply_contract.number }}")
	_contract_centered(doc, "от {{ supply_contract_date_fmt }} г.")
	_contract_body_para(
		doc,
		"{% if seller.company_type != 'ИП' %}"
		"{{ seller.company_name }}, далее «Поставщик», в лице {{ supply_contract.officials[0].position or '_____________' }} "
		"{% if supply_contract.officials[0].full_name %}{{ supply_contract.officials[0].full_name }}{% else %}_____________{% endif %}, "
		"{% if supply_contract.officials[0].is_base %}действующего на основании {{ supply_contract.officials[0].base_document }} "
		"{{ supply_contract.officials[0].base_document_name }}{% endif %}, с одной стороны, "
		"{% else %}"
		"{{ seller.company_name }}, далее «Поставщик», ОГРНИП {{ seller.ogrn }}, с одной стороны, "
		"{% endif %}",
	)
	_contract_body_para(
		doc,
		"{% if buyer.company_type != 'ИП' %}"
		"{{ buyer.company_name }}, далее «Покупатель», с другой стороны, "
		"{% else %}"
		"{{ buyer.company_name }}, далее «Покупатель», ОГРНИП {{ buyer.ogrn }}, с другой стороны, "
		"{% endif %}"
		"заключили настоящий Договор поставки о нижеследующем:",
	)
	_contract_body_para(
		doc,
		"{% if supply_contract.supply_contract_text %}{{ supply_contract.supply_contract_text }}{% endif %}",
	)
	_contract_body_para(doc, "Спецификация № {{ supply_contract.specification_number }} от {{ specification_date_fmt }} г.")
	_contract_body_para(
		doc,
		"{% for item in items %}"
		"{{ loop.index }}. {{ item.product_name }} — {{ item.quantity }} {{ item.unit_of_measurement }} × "
		"{{ item.price }} = {{ item.amount }}; "
		"{% endfor %}",
	)
	_contract_body_para(
		doc,
		"Сумма без НДС: {{ total_amount_excl_vat }}, НДС: {{ amount_vat_rate }}, итого: {{ total_amount }}",
	)
	_contract_body_para(
		doc,
		"{% if supply_contract.specification_text %}{{ supply_contract.specification_text }}{% endif %}",
	)
	_add_two_column_lines(
		doc,
		[
			(
				"{% if supply_contract.supplier_details_check %}ПОСТАВЩИК:{% endif %}",
				"{% if supply_contract.buyer_details_check %}ПОКУПАТЕЛЬ:{% endif %}",
			),
			(
				"{% if supply_contract.supplier_details_check %}{{ seller.company_name }}{% endif %}",
				"{% if supply_contract.buyer_details_check %}{{ buyer.company_name }}{% endif %}",
			),
			(
				"{% if supply_contract.supplier_details_check %}{{ seller.index }}, {{ seller.legal_address }}{% endif %}",
				"{% if supply_contract.buyer_details_check %}{{ buyer.index }}, {{ buyer.legal_address }}{% endif %}",
			),
			(
				"{% if supply_contract.supplier_details_check %}ИНН {{ seller.inn }}{% endif %}",
				"{% if supply_contract.buyer_details_check %}ИНН {{ buyer.inn }}{% endif %}",
			),
			(
				"{% if supply_contract.supplier_details_check %}КПП {{ seller.kpp }}{% endif %}",
				"{% if supply_contract.buyer_details_check %}КПП {{ buyer.kpp }}{% endif %}",
			),
			(
				"{% if supply_contract.supplier_details_check %}Рас/счет № {{ seller.account_number or '' }} в {{ seller.bank_name or '' }}{% endif %}",
				"{% if supply_contract.buyer_details_check %}Рас/счет № {{ buyer.account_number or '' }} в {{ buyer.bank_name or '' }}{% endif %}",
			),
			(
				"{% if supply_contract.supplier_details_check %}{{ seller.correspondent_bank_account or '' }}{% endif %}",
				"{% if supply_contract.buyer_details_check %}{{ buyer.correspondent_bank_account or '' }}{% endif %}",
			),
			(
				"{% if supply_contract.supplier_details_check %}{{ seller.bic or '' }}{% endif %}",
				"{% if supply_contract.buyer_details_check %}{{ buyer.bic or '' }}{% endif %}",
			),
			(
				"{% if supply_contract.supplier_details_check %}{{ seller.email }}{% endif %}",
				"{% if supply_contract.buyer_details_check %}{{ buyer.email }}{% endif %}",
			),
			(
				"{% if supply_contract.supplier_details_check %}{{ seller.phone }}{% endif %}",
				"{% if supply_contract.buyer_details_check %}{{ buyer.phone }}{% endif %}",
			),
		],
		bold_header=True,
		page_break_before=True,
		keep_rows_together=True,
	)
	_add_two_column_lines(
		doc,
		[
			("Поставщик:", "Покупатель:"),
			("{{ seller.company_name }}", "{{ buyer.company_name }}"),
			(
				"{{ supply_contract.officials[0].position or '_________________(ДОЛЖНОСТЬ)' }}",
				"_________________(ДОЛЖНОСТЬ)",
			),
			("", ""),
			("_________________", "_________________"),
			(
				"/{% if supply_contract.officials[0].full_name %}{{ supply_contract.officials[0].full_name }}{% else %}_____________(ФИО){% endif %}/",
				"/{% if buyer.full_name %}{{ buyer.full_name }}{% else %}_____________(ФИО){% endif %}/",
			),
			("«____» _______________ 20__г.", "«____» _______________ 20__г."),
			("М.П.", "М.П."),
		],
		bold_header=True,
		keep_rows_together=True,
	)
	path.parent.mkdir(parents=True, exist_ok=True)
	doc.save(str(path))
	print("Wrote", path)


def main() -> None:
	root = Path(__file__).resolve().parents[1]
	docx_dir = root / "app" / "templates" / "docx"
	write_bill_payment(docx_dir / "bill.docx")
	write_bill_contract(docx_dir / "bill_contract.docx")
	write_bill_offer(docx_dir / "bill_offer.docx")
	write_supply_contract(docx_dir / "supply_contract.docx")


if __name__ == "__main__":
	main()
